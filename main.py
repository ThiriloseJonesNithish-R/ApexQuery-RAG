import os
import sys
import glob
from dotenv import load_dotenv
from colorama import init, Fore, Style

# Initialize colorama
init(autoreset=True)

# Load env variables from .env file
load_dotenv()

def print_banner():
    banner = f"""
{Fore.CYAN}{Style.BRIGHT}=============================================================
             AI DOCUMENT QUERY SYSTEM (RAG App)
=============================================================
{Fore.WHITE}Analyze PDFs, Word files, and Text documents instantly.
Ask questions based directly on their contents.
{Fore.CYAN}============================================================={Style.RESET_ALL}
"""
    print(banner)

def extract_txt(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    except Exception as e:
        print(f"{Fore.RED}Error reading text file {file_path}: {e}")
        return ""

def extract_pdf(file_path):
    try:
        from pypdf import PdfReader
        reader = PdfReader(file_path)
        text = []
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text.append(page_text)
        return "\n".join(text)
    except ImportError:
        print(f"{Fore.RED}Error: 'pypdf' package is not installed.")
        return ""
    except Exception as e:
        print(f"{Fore.RED}Error reading PDF file {file_path}: {e}")
        return ""

def extract_docx(file_path):
    try:
        import docx
        doc = docx.Document(file_path)
        text = []
        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text)
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text.append(" | ".join(row_text))
        return "\n".join(text)
    except ImportError:
        print(f"{Fore.RED}Error: 'python-docx' package is not installed.")
        return ""
    except Exception as e:
        print(f"{Fore.RED}Error reading Word file {file_path}: {e}")
        return ""

def process_file(file_path):
    if not os.path.exists(file_path):
        print(f"{Fore.RED}File not found: {file_path}")
        return None
    
    ext = os.path.splitext(file_path)[1].lower()
    print(f"{Fore.YELLOW}Processing {os.path.basename(file_path)}...")
    
    if ext == '.txt':
        content = extract_txt(file_path)
    elif ext == '.pdf':
        content = extract_pdf(file_path)
    elif ext in ['.docx', '.doc']:
        content = extract_docx(file_path)
    else:
        print(f"{Fore.RED}Unsupported file extension '{ext}'. Only .txt, .pdf, and .docx are supported.")
        return None
        
    if not content or not content.strip():
        print(f"{Fore.RED}Warning: Extracted text from {file_path} is empty.")
        return ""
    
    word_count = len(content.split())
    print(f"{Fore.GREEN}Successfully parsed {os.path.basename(file_path)} (~{word_count} words).")
    return content

def get_gemini_client():
    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        print(f"{Fore.RED}Error: GEMINI_API_KEY environment variable is not set.")
        print(f"{Fore.YELLOW}Please create a '.env' file in this directory and add: GEMINI_API_KEY=your_api_key_here")
        sys.exit(1)
        
    try:
        from google import genai
        # genai.Client automatically uses GEMINI_API_KEY from env
        return genai.Client()
    except ImportError:
        print(f"{Fore.RED}Error: 'google-genai' SDK is not installed.")
        print(f"{Fore.YELLOW}Please run: pip install google-genai")
        sys.exit(1)
    except Exception as e:
        print(f"{Fore.RED}Failed to initialize Gemini Client: {e}")
        sys.exit(1)

def main():
    print_banner()
    
    # 1. Ask for files to load
    files_to_load = []
    if len(sys.argv) > 1:
        files_to_load = sys.argv[1:]
    else:
        print(f"{Fore.CYAN}No files specified in arguments. Let's add them now.")
        user_input = input(f"{Fore.WHITE}Enter file path(s) (comma-separated, or use wildcards like *.pdf): {Style.RESET_ALL}").strip()
        if not user_input:
            print(f"{Fore.RED}No files provided. Exiting.")
            return
            
        # Parse inputs, handle glob and comma separation
        parts = [p.strip() for p in user_input.split(',')]
        for part in parts:
            matched = glob.glob(part)
            if matched:
                files_to_load.extend(matched)
            else:
                files_to_load.append(part)

    # 2. Extract contents
    combined_content = []
    for file_path in files_to_load:
        content = process_file(file_path)
        if content:
            combined_content.append(f"--- START OF FILE: {os.path.basename(file_path)} ---\n{content}\n--- END OF FILE: {os.path.basename(file_path)} ---")
            
    if not combined_content:
        print(f"{Fore.RED}No valid documents were successfully loaded. Exiting.")
        return
        
    full_document_text = "\n\n".join(combined_content)
    total_words = len(full_document_text.split())
    print(f"\n{Fore.GREEN}{Style.BRIGHT}All documents loaded successfully! Total text size: ~{total_words} words.")
    
    # 3. Setup Gemini Chat
    print(f"{Fore.YELLOW}Initializing AI model Chat Session...")
    client = get_gemini_client()
    
    from google.genai import types
    
    system_instruction = (
        "You are a helpful, precise AI assistant that answers questions based on the uploaded document contents.\n"
        "Here are the contents of the document(s) uploaded by the user:\n\n"
        f"{full_document_text}\n\n"
        "Provide accurate answers referencing the context where possible. If the answer cannot be found or inferred from the text, "
        "state that it is not present in the documents, but try to answer as best as you can if there's related information."
    )
    
    try:
        chat = client.chats.create(
            model='gemini-3.5-flash',
            config=types.GenerateContentConfig(
                system_instruction=system_instruction,
                temperature=0.2, # Lower temperature for factual accuracy
            )
        )
    except Exception as e:
        print(f"{Fore.RED}Error starting chat session: {e}")
        return
        
    print(f"{Fore.GREEN}AI Chat Session is ready!{Style.RESET_ALL}")
    print(f"{Fore.CYAN}Type your questions below. Type {Fore.YELLOW}'exit'{Fore.CYAN} or {Fore.YELLOW}'quit'{Fore.CYAN} to end the session.\n")
    
    # 4. Interactive Q&A loop
    while True:
        try:
            question = input(f"{Fore.BLUE}{Style.BRIGHT}User > {Style.RESET_ALL}").strip()
            if not question:
                continue
            if question.lower() in ['exit', 'quit']:
                print(f"\n{Fore.CYAN}Ending session. Goodbye!")
                break
                
            print(f"{Fore.GREEN}{Style.BRIGHT}AI > {Style.RESET_ALL}", end="", flush=True)
            
            # Send message and stream the response
            response_stream = chat.send_message_stream(question)
            for chunk in response_stream:
                print(chunk.text, end="", flush=True)
            print("\n") # New line after stream ends
            
        except KeyboardInterrupt:
            print(f"\n{Fore.CYAN}Session interrupted. Goodbye!")
            break
        except Exception as e:
            print(f"\n{Fore.RED}An error occurred during Q&A: {e}\n")

if __name__ == "__main__":
    main()
